"""Invoice extraction and pallet planning helpers for the packaging planner."""

from __future__ import annotations

import csv
import io
import math
import os
import re
import uuid
from collections import defaultdict


ITEM_TYPES = (
    "complete_table",
    "body_only",
    "top_rail_only",
    "cushion_only",
    "legs_only",
    "replacement_item",
    "other",
)

ITEM_TYPE_LABELS = {
    "complete_table": "Complete table",
    "body_only": "Body only",
    "top_rail_only": "Top rail only",
    "cushion_only": "Cushion set only",
    "legs_only": "Leg boxes only",
    "replacement_item": "Replacement item",
    "other": "Other component",
}

SUPPORTED_EXTENSIONS = {
    ".pdf", ".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff",
    ".csv", ".xlsx", ".xlsm", ".docx", ".txt",
}

MODEL_PATTERNS = (
    ("Premium Edition", re.compile(r"\bpremium(?:\s+edition)?\b", re.I)),
    ("Champion", re.compile(r"\bchampion\b", re.I)),
    ("League", re.compile(r"\bleague\b", re.I)),
)

COLOUR_PATTERNS = (
    ("Rustic Black", re.compile(r"\brustic\s+black\b", re.I)),
    ("Rustic Oak", re.compile(r"\brustic\s+oak\b", re.I)),
    ("Grey Oak", re.compile(r"\bgr[ae]y\s+oak\b", re.I)),
    ("Black", re.compile(r"\bblack\b", re.I)),
    ("Stone", re.compile(r"\bstone\b", re.I)),
    ("White", re.compile(r"\bwhite\b", re.I)),
)

FIELD_ALIASES = {
    "description": {
        "description", "product", "product description", "item", "item description",
        "product name", "details", "sku description",
    },
    "size": {"size", "table size", "feet", "ft"},
    "model": {"model", "range", "table model"},
    "colour": {"colour", "color", "finish"},
    "quantity": {"quantity", "qty", "order qty", "ordered", "units"},
    "po_number": {
        "po", "po number", "po no", "purchase order", "purchase order number",
        "order number", "customer po",
    },
    "item_type": {"type", "item type", "component", "product type"},
    "notes": {"notes", "note", "comments"},
}


def _new_id(prefix):
    return f"{prefix}-{uuid.uuid4().hex[:12]}"


def _clean_text(value):
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def _safe_quantity(value, default=1):
    if value is None or value == "":
        return default
    match = re.search(r"-?\d+", str(value).replace(",", ""))
    if not match:
        return default
    return max(1, min(9999, int(match.group(0))))


def _normalise_header(value):
    return re.sub(r"[^a-z0-9]+", " ", _clean_text(value).lower()).strip()


def _field_for_header(header):
    normalised = _normalise_header(header)
    for field_name, aliases in FIELD_ALIASES.items():
        if normalised in aliases:
            return field_name
    return None


def infer_size(text):
    text = _clean_text(text)
    if re.search(r"\b6\s*(?:ft|foot|feet|')\b", text, re.I):
        return "6ft"
    if re.search(r"\b7\s*(?:ft|foot|feet|')\b", text, re.I):
        return "7ft"
    return ""


def infer_model(text):
    for label, pattern in MODEL_PATTERNS:
        if pattern.search(text or ""):
            return label
    return ""


def infer_colour(text):
    for label, pattern in COLOUR_PATTERNS:
        if pattern.search(text or ""):
            return label
    return ""


def infer_item_type(text):
    lowered = _clean_text(text).lower()
    if ("top rail" in lowered or "toprail" in lowered) and (
        "replacement" in lowered or "only" in lowered
    ):
        return "top_rail_only"
    if "top rail" in lowered or "toprail" in lowered:
        return "top_rail_only"
    if "cushion" in lowered:
        return "cushion_only"
    if "body only" in lowered or "table body" in lowered or lowered.startswith("body "):
        return "body_only"
    if "leg box" in lowered or "legs only" in lowered:
        return "legs_only"
    if "replacement" in lowered or "spare" in lowered:
        return "replacement_item"
    if (
        "complete table" in lowered
        or "pool table" in lowered
        or any(pattern.search(lowered) for _, pattern in MODEL_PATTERNS)
    ):
        return "complete_table"
    return "other"


def normalise_item_type(value, description=""):
    value_text = _clean_text(value).lower().replace("-", " ").replace("_", " ")
    direct = {
        "complete": "complete_table",
        "complete table": "complete_table",
        "table": "complete_table",
        "body": "body_only",
        "body only": "body_only",
        "table body": "body_only",
        "top rail": "top_rail_only",
        "top rail only": "top_rail_only",
        "replacement top rail": "top_rail_only",
        "cushion": "cushion_only",
        "cushions": "cushion_only",
        "cushion set": "cushion_only",
        "cushion set only": "cushion_only",
        "legs": "legs_only",
        "leg boxes": "legs_only",
        "replacement": "replacement_item",
        "replacement item": "replacement_item",
        "other": "other",
    }
    if value_text in direct:
        return direct[value_text]
    underscored = value_text.replace(" ", "_")
    if underscored in ITEM_TYPES:
        return underscored
    return infer_item_type(description)


def build_item(values, source_file="", raw_text="", confidence=None):
    description = _clean_text(values.get("description") or raw_text)
    combined = " ".join(
        filter(None, [
            description,
            _clean_text(values.get("size")),
            _clean_text(values.get("model")),
            _clean_text(values.get("colour")),
            _clean_text(values.get("item_type")),
        ])
    )
    size = infer_size(values.get("size") or combined)
    model = _clean_text(values.get("model")) or infer_model(combined)
    colour = _clean_text(values.get("colour")) or infer_colour(combined)
    item_type = normalise_item_type(values.get("item_type"), description)
    quantity = _safe_quantity(values.get("quantity"), 1)
    po_number = _clean_text(values.get("po_number"))
    notes = _clean_text(values.get("notes"))

    if confidence is None:
        recognised = sum(bool(value) for value in (size, model, colour))
        confidence = 0.45 + (0.12 * recognised)
        if item_type != "other":
            confidence += 0.15
        if po_number:
            confidence += 0.05
        confidence = min(confidence, 0.98)

    return {
        "id": _new_id("item"),
        "source_file": _clean_text(source_file),
        "description": description or "Unrecognised invoice item",
        "size": size,
        "model": model,
        "colour": colour,
        "quantity": quantity,
        "item_type": item_type,
        "po_number": po_number,
        "notes": notes,
        "confidence": round(float(confidence), 2),
        "raw_text": _clean_text(raw_text)[:2000],
    }


def _extract_po(text):
    compact_match = re.search(r"\b(PO[A-Z0-9][A-Z0-9./_-]{2,})\b", text or "", re.I)
    if compact_match:
        return compact_match.group(1).upper()

    patterns = (
        r"\b(?:purchase\s+order|customer\s+po|po(?:\s+number|\s+no\.?)?)\s*[:#-]?\s*([A-Z0-9][A-Z0-9./_-]{2,})",
        r"\border\s+(?:number|no\.?)\s*[:#-]?\s*([A-Z0-9][A-Z0-9./_-]{2,})",
    )
    for pattern in patterns:
        match = re.search(pattern, text or "", re.I)
        if match:
            return match.group(1)
    return ""


def _extract_document_po(text):
    for raw_line in (text or "").splitlines():
        line = _clean_text(raw_line)
        if re.match(
            r"^(?:purchase\s+order|customer\s+po|po\s+(?:number|no\.?))\s*[:#-]",
            line,
            re.I,
        ):
            return _extract_po(line)
    return ""


def _line_quantity(line):
    match = re.match(r"^\s*(\d+)\s*(?:x|X|\*)\s+", line)
    if match:
        return int(match.group(1))
    match = re.search(r"\bqty\.?\s*[:x-]?\s*(\d+)\b", line, re.I)
    return int(match.group(1)) if match else 1


def _following_quantity(lines, product_index):
    for next_line in lines[product_index + 1:product_index + 4]:
        match = re.match(
            r"^\s*(\d+)\s+units?\s*(?:x|X|\*)\s+",
            next_line,
            re.I,
        )
        if match:
            return max(1, int(match.group(1)))
        if next_line and not re.match(r"^[^\w]*$", next_line):
            break
    return None


def _clean_product_description(line):
    return re.sub(
        r"\s+(?:\u00a3|\u0141|\$|\u20ac)\s*[\d,]+(?:\.\d{2})?\s*$",
        "",
        line,
    ).strip()


def items_from_text(text, source_file):
    text = (text or "").replace("\x00", " ")
    global_po = _extract_document_po(text)
    candidates = []
    seen = set()
    product_terms = re.compile(
        r"\b(?:6\s*ft|7\s*ft|champion|league|premium|pool\s+table|table\s+body|"
        r"top\s*rail|cushion|leg\s*box|replacement)\b",
        re.I,
    )
    ignored_terms = re.compile(
        r"^(?:invoice|description|product|quantity|qty|subtotal|total|vat|delivery|"
        r"purchase order|page \d+)\b",
        re.I,
    )
    lines = [_clean_text(raw_line) for raw_line in text.splitlines()]
    for line_index, line in enumerate(lines):
        if len(line) < 4 or ignored_terms.search(line):
            continue
        if not product_terms.search(line):
            continue
        fingerprint = line.lower()
        if fingerprint in seen:
            continue
        seen.add(fingerprint)
        following_quantity = _following_quantity(lines, line_index)
        description = _clean_product_description(line)
        candidates.append(build_item({
            "description": description,
            "quantity": (
                following_quantity
                if following_quantity is not None
                else _line_quantity(line)
            ),
            "po_number": _extract_po(line) or global_po,
        }, source_file=source_file, raw_text=line))
    return candidates


def items_from_rows(rows, source_file):
    rows = [list(row) for row in rows if any(_clean_text(cell) for cell in row)]
    if not rows:
        return []

    header_index = None
    field_map = {}
    for index, row in enumerate(rows[:20]):
        candidate_map = {
            column_index: _field_for_header(cell)
            for column_index, cell in enumerate(row)
            if _field_for_header(cell)
        }
        if "description" in candidate_map.values() or len(candidate_map) >= 3:
            header_index = index
            field_map = candidate_map
            break

    if header_index is None:
        return items_from_text(
            "\n".join(" | ".join(_clean_text(cell) for cell in row) for row in rows),
            source_file,
        )

    items = []
    global_po = ""
    for row in rows[header_index + 1:]:
        values = {}
        for column_index, field_name in field_map.items():
            if column_index < len(row):
                values[field_name] = _clean_text(row[column_index])
        if values.get("po_number"):
            global_po = values["po_number"]
        elif global_po:
            values["po_number"] = global_po
        description = values.get("description", "")
        if not description or _normalise_header(description) in FIELD_ALIASES["description"]:
            continue
        items.append(build_item(values, source_file=source_file, raw_text=" | ".join(
            _clean_text(cell) for cell in row
        ), confidence=0.9))
    return items


def _decode_text(data):
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_invoice_bytes(filename, data):
    extension = os.path.splitext(filename or "")[1].lower()
    warnings = []
    items = []

    try:
        if extension == ".csv":
            text = _decode_text(data)
            sample = text[:4096]
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
            except csv.Error:
                dialect = csv.excel
            rows = list(csv.reader(io.StringIO(text), dialect))
            items = items_from_rows(rows, filename)
        elif extension in (".xlsx", ".xlsm"):
            try:
                import openpyxl
            except ImportError:
                warnings.append(
                    f"{filename}: Excel support is not installed. Add the item manually."
                )
            else:
                workbook = openpyxl.load_workbook(
                    io.BytesIO(data), read_only=True, data_only=True
                )
                for sheet in workbook.worksheets:
                    items.extend(items_from_rows(sheet.iter_rows(values_only=True), filename))
        elif extension == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError:
                warnings.append(
                    f"{filename}: PDF support is not installed. Add the item manually."
                )
            else:
                reader = PdfReader(io.BytesIO(data))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
                items = items_from_text(text, filename)
                if not text.strip():
                    warnings.append(
                        f"{filename}: no selectable PDF text was found. "
                        "It may be a scanned invoice; review it manually."
                    )
        elif extension == ".docx":
            try:
                from docx import Document
            except ImportError:
                warnings.append(
                    f"{filename}: Word support is not installed. Add the item manually."
                )
            else:
                document = Document(io.BytesIO(data))
                text_parts = [paragraph.text for paragraph in document.paragraphs]
                for table in document.tables:
                    rows = [[cell.text for cell in row.cells] for row in table.rows]
                    items.extend(items_from_rows(rows, filename))
                items.extend(items_from_text("\n".join(text_parts), filename))
        elif extension in (".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"):
            try:
                import pytesseract
                from PIL import Image
            except ImportError:
                warnings.append(
                    f"{filename}: image OCR is not installed. Add the item manually."
                )
            else:
                try:
                    text = pytesseract.image_to_string(Image.open(io.BytesIO(data)))
                    items = items_from_text(text, filename)
                except Exception:
                    warnings.append(
                        f"{filename}: image OCR was unavailable or could not read the file. "
                        "Review it manually."
                    )
        elif extension == ".txt":
            items = items_from_text(_decode_text(data), filename)
        else:
            warnings.append(f"{filename}: unsupported file type.")
    except Exception as error:
        warnings.append(f"{filename}: extraction failed ({type(error).__name__}).")

    if not items:
        warnings.append(
            f"{filename}: no products were confidently detected. "
            "Use Add item to enter the invoice lines."
        )
    return items, warnings


def extract_invoice_files(files):
    all_items = []
    warnings = []
    source_files = []
    for uploaded_file in files:
        filename = os.path.basename(uploaded_file.filename or "invoice")
        source_files.append(filename)
        data = uploaded_file.read()
        items, file_warnings = extract_invoice_bytes(filename, data)
        all_items.extend(items)
        warnings.extend(file_warnings)
    return all_items, warnings, source_files


def normalise_items(items):
    clean_items = []
    for item in items or []:
        item = dict(item or {})
        clean = build_item(item, source_file=item.get("source_file", ""))
        clean["id"] = _clean_text(item.get("id")) or clean["id"]
        clean["confidence"] = max(0.0, min(1.0, float(item.get("confidence", clean["confidence"]))))
        clean["raw_text"] = _clean_text(item.get("raw_text"))[:2000]
        clean_items.append(clean)
    return clean_items


def normalise_config(config):
    config = dict(config or {})
    return {
        "body_capacity": max(1, min(50, _safe_quantity(config.get("body_capacity"), 5))),
        "top_rail_capacity": max(1, min(100, _safe_quantity(config.get("top_rail_capacity"), 15))),
        "loose_rail_limit": max(1, min(100, _safe_quantity(config.get("loose_rail_limit"), 10))),
        "top_rails_per_body_pallet": max(
            1,
            min(20, _safe_quantity(config.get("top_rails_per_body_pallet"), 2)),
        ),
        "cushion_pallet_count": max(1, min(50, _safe_quantity(config.get("cushion_pallet_count"), 1))),
        "legs_per_6ft_table": max(0, min(100, int(config.get("legs_per_6ft_table", 4) or 0))),
        "legs_per_box": max(1, min(100, _safe_quantity(config.get("legs_per_box"), 8))),
    }


def _component_line(item, component_type):
    return {
        "id": _new_id("line"),
        "item_id": item["id"],
        "component_type": component_type,
        "size": item.get("size", ""),
        "model": item.get("model", ""),
        "colour": item.get("colour", ""),
        "quantity": item["quantity"],
        "po_number": item.get("po_number", ""),
        "description": item.get("description", ""),
        "notes": item.get("notes", ""),
        "source_file": item.get("source_file", ""),
        "origin_type": item.get("item_type", ""),
    }


def build_requirements(items, config):
    requirements = {
        "bodies": [],
        "top_rails": [],
        "cushions": [],
        "leg_boxes": [],
        "complete_6ft": 0,
        "complete_tables": 0,
    }
    for item in items:
        item_type = item["item_type"]
        if item_type == "complete_table":
            requirements["complete_tables"] += item["quantity"]
            if item["size"] == "6ft":
                requirements["complete_6ft"] += item["quantity"]
            requirements["bodies"].append(_component_line(item, "body"))
            requirements["top_rails"].append(_component_line(item, "top_rail"))
            requirements["cushions"].append(_component_line(item, "cushion"))
        elif item_type == "body_only":
            requirements["bodies"].append(_component_line(item, "body"))
        elif item_type == "top_rail_only":
            requirements["top_rails"].append(_component_line(item, "top_rail"))
        elif item_type == "cushion_only":
            requirements["cushions"].append(_component_line(item, "cushion"))
        elif item_type == "legs_only":
            requirements["leg_boxes"].append(_component_line(item, "leg_box"))
        elif item_type == "replacement_item":
            inferred = infer_item_type(item.get("description", ""))
            if inferred == "body_only":
                requirements["bodies"].append(_component_line(item, "body"))
            elif inferred == "top_rail_only":
                requirements["top_rails"].append(_component_line(item, "top_rail"))
            elif inferred == "cushion_only":
                requirements["cushions"].append(_component_line(item, "cushion"))

    required_legs = requirements["complete_6ft"] * config["legs_per_6ft_table"]
    automatic_leg_boxes = math.ceil(required_legs / config["legs_per_box"]) if required_legs else 0
    if automatic_leg_boxes:
        requirements["leg_boxes"].append({
            "id": _new_id("line"),
            "item_id": "",
            "component_type": "leg_box",
            "size": "6ft",
            "model": "",
            "colour": "",
            "quantity": automatic_leg_boxes,
            "po_number": "",
            "description": "Leg boxes for complete 6ft tables",
            "notes": (
                f"{required_legs} legs required; "
                f"{config['legs_per_box']} legs per box"
            ),
            "source_file": "",
            "origin_type": "automatic",
        })
    requirements["required_legs"] = required_legs
    requirements["automatic_leg_boxes"] = automatic_leg_boxes
    return requirements


def _take_from_lines(lines, quantity):
    taken = []
    remaining = quantity
    while lines and remaining > 0:
        line = lines[0]
        amount = min(line["quantity"], remaining)
        allocated = dict(line)
        allocated["id"] = _new_id("line")
        allocated["quantity"] = amount
        taken.append(allocated)
        line["quantity"] -= amount
        remaining -= amount
        if line["quantity"] <= 0:
            lines.pop(0)
    return taken


def _new_pallet(pallet_type, number, lines=None, size=""):
    return {
        "id": _new_id("pallet"),
        "pallet_number": number,
        "pallet_type": pallet_type,
        "size": size,
        "is_mixed": False,
        "manual_override": False,
        "notes": "",
        "lines": lines or [],
        "carried_top_rails": [],
    }


def _line_total(lines):
    return sum(max(0, int(line.get("quantity", 0) or 0)) for line in lines or [])


def _aggregate_cushion_lines(lines):
    grouped = defaultdict(list)
    for line in lines:
        grouped[line.get("size") or "Unknown"].append(line)

    aggregated = []
    for size in ("6ft", "7ft", "Unknown"):
        size_lines = grouped.get(size, [])
        quantity = _line_total(size_lines)
        if not quantity:
            continue
        po_numbers = sorted({
            line.get("po_number")
            for line in size_lines
            if line.get("po_number")
        })
        aggregated.append({
            "id": _new_id("line"),
            "item_id": "",
            "component_type": "cushion",
            "size": "" if size == "Unknown" else size,
            "model": "",
            "colour": "",
            "quantity": quantity,
            "po_number": ", ".join(po_numbers),
            "description": (
                f"All {size} cushion sets"
                if size != "Unknown"
                else "Cushion sets with size to confirm"
            ),
            "notes": "Combined cushion total from all invoice items",
            "source_file": "",
            "origin_type": "automatic",
        })
    return aggregated


def _refresh_pallet_labels(pallet, config):
    body_sizes = {
        line.get("size") for line in pallet.get("lines", [])
        if line.get("component_type") == "body" and line.get("size")
    }
    all_sizes = {
        line.get("size") for line in pallet.get("lines", [])
        if line.get("size")
    }
    pallet["is_mixed"] = bool(pallet.get("is_mixed")) or len(all_sizes) > 1
    if pallet["pallet_type"] == "body":
        total = _line_total(pallet["lines"])
        pallet["capacity"] = config["body_capacity"]
        pallet["capacity_used"] = total
        pallet["status"] = "full" if total == config["body_capacity"] else "partial"
        if len(body_sizes) > 1:
            pallet["status"] = "mixed"
    elif pallet["pallet_type"] == "top_rail":
        total = _line_total(pallet["lines"])
        pallet["capacity"] = config["top_rail_capacity"]
        pallet["capacity_used"] = total
        pallet["status"] = "full" if total == config["top_rail_capacity"] else "partial"
    else:
        pallet["capacity"] = None
        pallet["capacity_used"] = _line_total(pallet["lines"])
        pallet["status"] = "mixed" if len(all_sizes) > 1 else "partial"
    return pallet


def _compatible_body_pallet(
    body_pallets,
    rail_size,
    rail_colour,
    config,
    required_rail_space=1,
):
    candidates = []
    for pallet in body_pallets:
        body_lines = [
            line for line in pallet["lines"]
            if line.get("component_type") == "body"
        ]
        sizes = {line.get("size") for line in body_lines}
        if rail_size == "7ft" and "7ft" not in sizes:
            continue
        if rail_size == "6ft" and not sizes.intersection({"6ft", "7ft"}):
            continue
        carried_lines = pallet.get("carried_top_rails", [])
        carried_count = _line_total(carried_lines)
        rail_capacity = max(
            config["top_rails_per_body_pallet"],
            _line_total(body_lines),
        )
        available_space = rail_capacity - carried_count
        if available_space < required_rail_space:
            continue
        candidates.append(pallet)
    candidates.sort(key=lambda pallet: (
        0 if rail_colour and rail_colour in {
            line.get("colour") for line in pallet.get("lines", [])
            if line.get("component_type") == "body" and line.get("colour")
        } else 1,
        _line_total(pallet.get("carried_top_rails", [])),
        pallet.get("capacity_used", 0) < pallet.get("capacity", 5),
        pallet["pallet_number"],
    ))
    return candidates[0] if candidates else None


def generate_packaging(items, config=None):
    items = normalise_items(items)
    config = normalise_config(config)
    requirements = build_requirements(items, config)
    pallets = []
    next_number = 1

    body_lines = [dict(line) for line in requirements["bodies"]]
    body_lines.sort(key=lambda line: (
        0 if line.get("size") == "7ft" else 1,
        line.get("colour") or "Unknown",
    ))
    while _line_total(body_lines):
        lines = _take_from_lines(body_lines, config["body_capacity"])
        pallet = _new_pallet("body", next_number, lines)
        _refresh_pallet_labels(pallet, config)
        pallets.append(pallet)
        next_number += 1
    body_pallets = list(pallets)

    rail_groups = defaultdict(list)
    for line in requirements["top_rails"]:
        group_key = (
            line.get("size") or "Unknown",
            line.get("colour") or "Unknown",
        )
        rail_groups[group_key].append(dict(line))
    for size in ("7ft", "6ft", "Unknown"):
        rail_batches = []
        colours = sorted(
            colour for group_size, colour in rail_groups
            if group_size == size
        )
        for colour in colours:
            rail_lines = rail_groups.get((size, colour), [])
            while _line_total(rail_lines) > config["top_rail_capacity"]:
                rail_batches.append(
                    _take_from_lines(rail_lines, config["top_rail_capacity"])
                )

            remainder = _line_total(rail_lines)
            if not remainder:
                continue
            rail_batches.append(_take_from_lines(rail_lines, remainder))

        planning_body_pallets = [
            {
                **pallet,
                "carried_top_rails": [
                    dict(line) for line in pallet.get("carried_top_rails", [])
                ],
            }
            for pallet in body_pallets
        ]
        planned_body_assignments = []
        all_batches_fit_bodies = (
            bool(rail_batches)
            and sum(_line_total(batch) for batch in rail_batches) < config["loose_rail_limit"]
        )
        for batch in rail_batches if all_batches_fit_bodies else []:
            batch_colour = next(
                (line.get("colour") for line in batch if line.get("colour")),
                "",
            )
            remaining_batch = [dict(line) for line in batch]
            while _line_total(remaining_batch):
                suitable_body = _compatible_body_pallet(
                    planning_body_pallets,
                    size,
                    batch_colour,
                    config,
                )
                if not suitable_body:
                    all_batches_fit_bodies = False
                    break
                body_count = _line_total([
                    line for line in suitable_body.get("lines", [])
                    if line.get("component_type") == "body"
                ])
                rail_capacity = max(config["top_rails_per_body_pallet"], body_count)
                available_space = rail_capacity - _line_total(
                    suitable_body.get("carried_top_rails", [])
                )
                moved_lines = _take_from_lines(
                    remaining_batch,
                    min(available_space, _line_total(remaining_batch)),
                )
                suitable_body["carried_top_rails"].extend(moved_lines)
                planned_body_assignments.append((suitable_body["id"], moved_lines))
            if not all_batches_fit_bodies:
                break

        if all_batches_fit_bodies:
            body_pallet_by_id = {pallet["id"]: pallet for pallet in body_pallets}
            for pallet_id, batch in planned_body_assignments:
                body_pallet_by_id[pallet_id]["carried_top_rails"].extend(batch)
            continue

        rail_batches.sort(key=_line_total, reverse=True)
        rail_pallet_lines = []
        for batch in rail_batches:
            batch_total = _line_total(batch)
            target_lines = next(
                (
                    lines for lines in rail_pallet_lines
                    if _line_total(lines) + batch_total <= config["top_rail_capacity"]
                ),
                None,
            )
            if target_lines is None:
                target_lines = []
                rail_pallet_lines.append(target_lines)
            target_lines.extend(batch)

        for lines in rail_pallet_lines:
            pallet = _new_pallet("top_rail", next_number, lines, size=size)
            _refresh_pallet_labels(pallet, config)
            pallets.append(pallet)
            next_number += 1

    cushion_lines = _aggregate_cushion_lines(requirements["cushions"])
    leg_lines = [dict(line) for line in requirements["leg_boxes"]]
    if _line_total(cushion_lines) or _line_total(leg_lines):
        cushion_pallet_count = config["cushion_pallet_count"]
        cushion_pallets = [
            _new_pallet("cushion", next_number + index)
            for index in range(cushion_pallet_count)
        ]
        next_number += cushion_pallet_count
        line_index = 0
        for source_lines in (cushion_lines, leg_lines):
            while source_lines:
                line = source_lines.pop(0)
                target = cushion_pallets[line_index % len(cushion_pallets)]
                target["lines"].append(line)
                line_index += 1
        for pallet in cushion_pallets:
            _refresh_pallet_labels(pallet, config)
            pallets.append(pallet)

    summary = build_summary(items, pallets, requirements, config)
    warnings = validate_packaging(items, pallets, config, requirements=requirements)
    return {
        "items": items,
        "pallets": pallets,
        "config": config,
        "summary": summary,
        "warnings": warnings,
    }


def build_summary(items, pallets, requirements=None, config=None):
    config = normalise_config(config)
    requirements = requirements or build_requirements(normalise_items(items), config)
    body_count = _line_total(requirements["bodies"])
    rail_count = _line_total(requirements["top_rails"])
    cushion_count = _line_total(requirements["cushions"])
    leg_box_count = _line_total(requirements["leg_boxes"])
    body_pallets = [pallet for pallet in pallets if pallet.get("pallet_type") == "body"]
    rail_pallets = [pallet for pallet in pallets if pallet.get("pallet_type") == "top_rail"]
    cushion_pallets = [pallet for pallet in pallets if pallet.get("pallet_type") == "cushion"]
    carried_rails = sum(
        _line_total(pallet.get("carried_top_rails", []))
        for pallet in body_pallets
    )

    cushion_by_size = defaultdict(int)
    for line in requirements["cushions"]:
        cushion_by_size[line.get("size") or "Unknown"] += line["quantity"]

    order_groups = defaultdict(int)
    for item in items:
        key = (
            item.get("size") or "Unknown",
            item.get("model") or "Unspecified",
            item.get("colour") or "Unspecified",
            ITEM_TYPE_LABELS.get(item.get("item_type"), "Other"),
            item.get("po_number") or "No PO",
        )
        order_groups[key] += item["quantity"]

    return {
        "complete_tables": requirements["complete_tables"],
        "total_bodies": body_count,
        "total_top_rails": rail_count,
        "total_cushion_sets": cushion_count,
        "total_leg_boxes": leg_box_count,
        "required_6ft_legs": requirements["required_legs"],
        "body_pallets": len(body_pallets),
        "top_rail_pallets": len(rail_pallets),
        "cushion_pallets": len(cushion_pallets),
        "physical_pallets": len(pallets),
        "carried_top_rails": carried_rails,
        "cushions_6ft": cushion_by_size["6ft"],
        "cushions_7ft": cushion_by_size["7ft"],
        "order_groups": [
            {
                "size": key[0],
                "model": key[1],
                "colour": key[2],
                "item_type": key[3],
                "po_number": key[4],
                "quantity": quantity,
            }
            for key, quantity in sorted(order_groups.items())
        ],
    }


def _warning(code, message, severity="warning"):
    return {
        "id": f"{code}-{uuid.uuid5(uuid.NAMESPACE_URL, message).hex[:10]}",
        "code": code,
        "message": message,
        "severity": severity,
    }


def validate_packaging(items, pallets, config=None, requirements=None):
    items = normalise_items(items)
    config = normalise_config(config)
    requirements = requirements or build_requirements(items, config)
    warnings = []

    required = {
        "body": _line_total(requirements["bodies"]),
        "top_rail": _line_total(requirements["top_rails"]),
        "cushion": _line_total(requirements["cushions"]),
        "leg_box": _line_total(requirements["leg_boxes"]),
    }
    packed = defaultdict(int)
    packed_by_size = defaultdict(int)
    for pallet in pallets or []:
        pallet_type = pallet.get("pallet_type")
        body_count = 0
        rail_count = 0
        for line in pallet.get("lines", []):
            component_type = line.get("component_type")
            quantity = max(0, int(line.get("quantity", 0) or 0))
            packed[component_type] += quantity
            packed_by_size[(component_type, line.get("size") or "Unknown")] += quantity
            if component_type == "body":
                body_count += quantity
            elif component_type == "top_rail":
                rail_count += quantity
            if (
                line.get("origin_type") == "replacement_item"
                and component_type in ("cushion", "leg_box")
            ):
                warnings.append(_warning(
                    "replacement_has_extras",
                    f"Replacement item '{line.get('description') or 'Unknown'}' has "
                    f"{component_type.replace('_', ' ')} contents assigned.",
                    "important",
                ))
        for line in pallet.get("carried_top_rails", []):
            quantity = max(0, int(line.get("quantity", 0) or 0))
            packed["top_rail"] += quantity
            packed_by_size[("top_rail", line.get("size") or "Unknown")] += quantity
            rail_count += quantity

        if pallet_type == "body" and body_count > config["body_capacity"]:
            warnings.append(_warning(
                "body_capacity",
                f"Pallet {pallet.get('pallet_number')} contains {body_count} bodies; "
                f"the maximum is {config['body_capacity']}.",
                "important",
            ))
        if pallet_type == "top_rail" and rail_count > config["top_rail_capacity"]:
            warnings.append(_warning(
                "rail_capacity",
                f"Pallet {pallet.get('pallet_number')} contains {rail_count} top rails; "
                f"the maximum is {config['top_rail_capacity']}.",
                "important",
            ))
        carried_count = _line_total(pallet.get("carried_top_rails", []))
        body_rail_capacity = max(config["top_rails_per_body_pallet"], body_count)
        if carried_count > body_rail_capacity:
            warnings.append(_warning(
                "body_pallet_rail_capacity",
                f"Pallet {pallet.get('pallet_number')} carries {carried_count} top rails; "
                f"the maximum on one body pallet is "
                f"{body_rail_capacity} (one per body).",
                "important",
            ))
        if carried_count >= config["loose_rail_limit"]:
            warnings.append(_warning(
                "too_many_carried_rails",
                f"Pallet {pallet.get('pallet_number')} carries {carried_count} loose top rails. "
                f"Use a dedicated pallet for {config['loose_rail_limit']} or more.",
                "important",
            ))
        if pallet_type == "body":
            body_sizes = {
                line.get("size") for line in pallet.get("lines", [])
                if line.get("component_type") == "body"
            }
            for line in pallet.get("carried_top_rails", []):
                if line.get("size") == "7ft" and "7ft" not in body_sizes:
                    warnings.append(_warning(
                        "unsafe_rail_assignment",
                        f"7ft top rails on pallet {pallet.get('pallet_number')} "
                        "do not have a 7ft body beneath them.",
                        "important",
                    ))
        actual_sizes = {
            line.get("size")
            for line in list(pallet.get("lines", [])) + list(pallet.get("carried_top_rails", []))
            if line.get("size")
        }
        if len(actual_sizes) > 1 and not pallet.get("is_mixed"):
            warnings.append(_warning(
                "unmarked_mixed_pallet",
                f"Pallet {pallet.get('pallet_number')} contains mixed sizes but is not marked mixed.",
            ))
        if not pallet.get("lines") and not pallet.get("carried_top_rails"):
            warnings.append(_warning(
                "empty_pallet",
                f"Pallet {pallet.get('pallet_number')} is empty.",
            ))

    labels = {
        "body": "table bodies",
        "top_rail": "top rails",
        "cushion": "cushion sets",
        "leg_box": "leg boxes",
    }
    for component_type, required_count in required.items():
        actual_count = packed[component_type]
        if actual_count != required_count:
            direction = "missing" if actual_count < required_count else "has extra"
            warnings.append(_warning(
                f"{component_type}_mismatch",
                f"Packaging {direction} {abs(required_count - actual_count)} "
                f"{labels[component_type]} (required {required_count}, packed {actual_count}).",
                "important",
            ))

    for component_type, requirement_lines in (
        ("body", requirements["bodies"]),
        ("top_rail", requirements["top_rails"]),
        ("cushion", requirements["cushions"]),
    ):
        required_sizes = defaultdict(int)
        for line in requirement_lines:
            required_sizes[line.get("size") or "Unknown"] += line["quantity"]
        all_sizes = set(required_sizes)
        all_sizes.update(
            size for packed_type, size in packed_by_size
            if packed_type == component_type
        )
        for size in sorted(all_sizes):
            expected = required_sizes[size]
            actual = packed_by_size[(component_type, size)]
            if expected != actual:
                warnings.append(_warning(
                    f"{component_type}_size_mismatch",
                    f"{size} {component_type.replace('_', ' ')} allocation mismatch: "
                    f"required {expected}, packed {actual}.",
                    "important",
                ))

    packed_leg_capacity = packed["leg_box"] * config["legs_per_box"]
    if packed_leg_capacity < requirements["required_legs"]:
        warnings.append(_warning(
            "insufficient_legs",
            f"{requirements['required_legs']} legs are required for the 6ft complete tables, "
            f"but the packed leg boxes hold only {packed_leg_capacity}.",
            "important",
        ))

    po_sources = defaultdict(set)
    for item in items:
        if item.get("po_number"):
            po_sources[item["po_number"].lower()].add(item.get("source_file") or "manual")
        if (
            item["item_type"] in (
                "complete_table", "body_only", "top_rail_only", "cushion_only"
            )
            and not item.get("size")
        ):
            warnings.append(_warning(
                "missing_table_size",
                f"'{item['description']}' has no 6ft or 7ft size.",
                "important",
            ))
        if item["item_type"] == "other":
            warnings.append(_warning(
                "unrecognised_item",
                f"Choose an order type for '{item['description']}' before using the final list.",
            ))
        if float(item.get("confidence", 1)) < 0.65:
            warnings.append(_warning(
                "low_confidence",
                f"Review '{item['description']}' from "
                f"{item.get('source_file') or 'manual entry'}; it was not confidently recognised.",
            ))
    for po_number, sources in po_sources.items():
        if len(sources) > 1:
            warnings.append(_warning(
                "duplicate_po",
                f"Purchase order {po_number.upper()} appears in multiple invoice files: "
                f"{', '.join(sorted(sources))}.",
            ))

    return warnings
